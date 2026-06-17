#!/usr/bin/env python3
"""Build a Chongqing regional talking-head operation package.

The script creates a Word execution manual and an Excel execution workbook for
one region. It intentionally generates a structured first draft, not verified
market research.
"""

from __future__ import annotations

import argparse
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


RISK_TERMS = [
    "必涨",
    "稳赚",
    "抄底",
    "闭眼买",
    "内部房源",
    "学区承诺",
    "制造焦虑",
    "零首付",
    "首付贷",
    "经营贷",
    "包成交",
    "捡漏",
    "大降价",
]


def safe_name(text: str) -> str:
    for ch in '\\/:*?"<>|':
        text = text.replace(ch, "")
    return text.strip() or "区域"


def check_risk(*parts: str) -> None:
    text = "\n".join(parts)
    found = [term for term in RISK_TERMS if term in text]
    if found:
        raise ValueError(f"存在高风险表达：{', '.join(found)}")


def style_sheet(ws, widths: dict[str, int] | None = None) -> None:
    widths = widths or {}
    header_fill = PatternFill("solid", fgColor="17365D")
    thin = Side(style="thin", color="D9E2F3")
    for cell in ws[1]:
        cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows():
        for cell in row:
            cell.font = Font(name="Microsoft YaHei", size=10, bold=cell.row == 1, color="FFFFFF" if cell.row == 1 else "000000")
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    ws.freeze_panes = "A2"


def add_rows(ws, rows: list[list[object]]) -> None:
    for row in rows:
        ws.append(row)


def research_rows(region: str) -> list[list[str]]:
    dimensions = [
        "市场大盘",
        "板块供需",
        "价格体系",
        "配套兑现",
        "交通通勤",
        "产品结构",
        "客群结构",
        "风险点",
    ]
    rows = []
    for i in range(40):
        dim = dimensions[i % len(dimensions)]
        rows.append([
            dim,
            f"{region}{dim}数据点{i + 1}",
            "待采集：正式版需填入可核验数值/事实",
            "待采集",
            "待采集：官方/平台/内部/第三方",
            "待补链接",
            "初稿占位，正式版需区分官方数据、平台参考价、第三方趋势和内部观察",
            f"用于{region}{dim}类选题",
            "待核验",
        ])
    return rows


def customer_rows(region: str) -> list[list[object]]:
    categories = ["还能不能买", "哪里能买", "适不适合我", "新房还是二手", "价格和流通", "配套兑现", "竞品板块对比"]
    rows = []
    for i in range(100):
        cat = categories[i % len(categories)]
        rows.append([i + 1, cat, f"客户关于{region}的{cat}问题样本{i + 1}，正式版替换为一线原始问题", "待采集：门店/带看/评论/私信", (i % 5) + 1])
    return rows


def competitor_rows(region: str) -> list[list[str]]:
    platforms = ["抖音", "小红书", "视频号", "B站", "知乎"]
    points = ["标题冲突", "前5秒问题", "数据证据", "评论追问", "风险替换"]
    rows = []
    for i in range(50):
        rows.append([
            platforms[i % len(platforms)],
            "待采集账号/发布方",
            f"{region}相关竞品内容方向{i + 1}",
            "待补公开链接",
            "待记录公开可见信息，不虚构点赞/播放",
            points[i % len(points)],
            "待拆解：观点-证据-结论-互动",
            "待记录评论区真实追问",
            "只借鉴结构，不复制原文",
            "替换高风险表达，避免夸张承诺",
        ])
    return rows


def topics(region: str) -> list[dict[str, str]]:
    base = [
        ("板块深挖", f"{region}现在到底处于什么阶段？", "数据型", "市场大盘+板块供需", "办公室地图墙"),
        ("板块深挖", f"{region}不是过时，而是进入筛选期", "竞品反向拆解", "竞品爆点+数据核验", "办公室+外景"),
        ("板块深挖", f"{region}核心区、边缘区怎么分？", "实地验证型", "轨道路径+商业半径", "地图墙"),
        ("配套交通", f"{region}商业到底成熟了吗？", "数据型", "商业开业/招商/人流观察", "商业外景"),
        ("配套交通", f"{region}轨道价值不能只看有没有站", "实地验证型", "站点步行路径", "轨道站口"),
        ("客户问题", f"150万预算，还能不能看{region}？", "客户问题型", "客户预算问题+平台参考价", "办公室图表"),
        ("客户问题", f"预算180到260万，{region}怎么选？", "客户问题型", "预算分层+产品结构", "办公室白板"),
        ("新房二手", f"{region}新房还是二手，怎么选？", "数据型", "新二手参考价+交付成本", "办公室图表"),
        ("价格流通", f"{region}挂牌多，是风险还是机会？", "数据型", "二手挂牌结构", "办公室数据板"),
        ("价格流通", f"{region}低总价房源，先问清这5件事", "客户问题型", "低总价原因排查", "办公室"),
        ("板块对比", f"{region}和礼嘉怎么比？", "竞品对比型", "板块对比+客户画像", "地图墙"),
        ("板块对比", f"{region}和悦来怎么选？", "竞品对比型", "兑现度+增量逻辑", "地图墙"),
        ("客户案例", f"为什么有客户从其他板块转看{region}？", "客户案例型", "匿名案例+需求排序", "办公室"),
        ("客户案例", f"{region}越看越乱，通常卡在哪里？", "客户问题型", "问题池高频母问题", "办公室"),
        ("总结清单", f"买{region}前，先做5个检查", "实地验证型", "数据+客户问题+竞品拆解", "办公室+地图墙"),
    ]
    full = []
    while len(full) < 30:
        item = base[len(full) % len(base)]
        full.append({
            "day": str(len(full) + 1),
            "column": item[0],
            "title": item[1] if len(full) < len(base) else f"{item[1]}（延展{len(full) + 1}）",
            "source": item[2],
            "evidence": item[3],
            "scene": item[4],
        })
    return full


def script_for_topic(region: str, item: dict[str, str]) -> str:
    return (
        f"结论先说：{item['title']}不能凭感觉判断，要先看数据、客户问题和现场验证。"
        f"第一，看证据锚点：{item['evidence']}，正式发布前要补齐时间、来源和口径。"
        f"第二，看客户适配：不同预算、通勤方向、家庭阶段，对{region}的判断不一样。"
        "第三，看风险边界：平台参考价不能说成成交价，规划和政策以官方当期口径为准。"
        f"所以这条视频不是简单说{region}好不好，而是帮客户知道自己该不该看、怎么看、先看哪里。"
    )


def build_workbook(region: str, account_line: str, out_path: Path, cover_dir: Path | None = None) -> None:
    wb = Workbook()
    wb.remove(wb.active)
    sheets = {
        "总览": ["模块", "数量/状态", "说明"],
        "数据资料包": ["维度", "数据点", "数值/事实", "时间", "来源名称", "链接", "口径", "可用于选题", "核验状态"],
        "客户问题池": ["序号", "分类", "原始客户问题", "来源", "频次"],
        "竞品拆解表": ["平台/来源", "账号/发布方", "样本标题/方向", "链接", "公开信息", "爆点类型", "内容结构拆解", "评论/客户追问", "可借鉴点", "风险替换"],
        "封面模板": ["模板", "标题结构", "适用栏目", "制作规则", "文件路径", "备注"],
        "30天选题排期": ["天", "栏目", "选题", "来源类型", "证据锚点", "拍摄场景"],
        "首批15条脚本卡": ["天", "栏目", "标题", "口播脚本", "封面主标题", "封面副标题", "视频号正文", "置顶评论", "证据锚点", "合规提示"],
        "拍摄通告单": ["序号", "拍摄时段", "脚本标题", "场景", "服装", "道具", "预计时长", "状态"],
        "发布复盘表": ["天", "发布日期", "发布时间", "播放", "完播率", "评论", "咨询", "有效线索", "下次优化"],
    }
    for name, headers in sheets.items():
        ws = wb.create_sheet(name)
        ws.append(headers)

    add_rows(wb["总览"], [
        ["数据资料包", "40条占位", "正式版不少于40条可核验数据"],
        ["客户问题池", "100条占位", "正式版替换为真实原始问题"],
        ["竞品拆解表", "50条占位", "只学习结构，不复制原文"],
        ["30天选题排期", "30条", "全部保留来源类型和证据锚点"],
        ["首批15条脚本卡", "15条", "初稿含证据提醒，正式版需补真实数据"],
        ["拍摄通告单", "15条", "按半月集中拍摄一天设计"],
        ["发布复盘表", "30条", "用于日更复盘"],
    ])
    add_rows(wb["数据资料包"], research_rows(region))
    add_rows(wb["客户问题池"], customer_rows(region))
    add_rows(wb["竞品拆解表"], competitor_rows(region))

    cover_rows = [
        ["数据判断型", f"{region}｜进入筛选期", "板块深挖/市场判断", "调用cq-koubo-fengmian，突出区域名+核心判断+3个数据标签", "", f"底部账号：{account_line}"],
        ["客户问题型", f"{region}｜220万怎么选", "预算/客户问题", "调用cq-koubo-fengmian，突出客户问题和决策清单", "", f"底部账号：{account_line}"],
        ["实地验证型", f"{region}｜核心区怎么分", "片区拆解/板块对比", "调用cq-koubo-fengmian，突出地图、路径和证据标签", "", f"底部账号：{account_line}"],
    ]
    if cover_dir and cover_dir.exists():
        for row in cover_rows:
            candidates = list(cover_dir.glob(f"*{row[0].split('型')[0]}*.png"))
            row[4] = str(candidates[0]) if candidates else str(cover_dir)
    add_rows(wb["封面模板"], cover_rows)

    items = topics(region)
    add_rows(wb["30天选题排期"], [[i["day"], i["column"], i["title"], i["source"], i["evidence"], i["scene"]] for i in items])
    script_rows = []
    for item in items[:15]:
        script_rows.append([
            item["day"],
            item["column"],
            item["title"],
            script_for_topic(region, item),
            region,
            item["title"].replace(region, "").strip("，？?"),
            f"{region}这类问题要结合数据、预算和真实生活半径判断。评论区留下预算和通勤方向，我按情况拆。",
            "你现在最纠结预算、通勤、配套，还是以后好不好转手？",
            item["evidence"],
            "平台参考价、挂牌价、成交价分开表达；政策和规划以官方当期口径为准。",
        ])
    add_rows(wb["首批15条脚本卡"], script_rows)
    add_rows(wb["拍摄通告单"], [[r[0], "上午/下午", r[2], items[int(r[0]) - 1]["scene"], "深色商务装", "地图/白板/提词器", "8-12分钟", "待拍摄"] for r in script_rows])
    add_rows(wb["发布复盘表"], [[i + 1, "", "20:00", "", "", "", "", "", ""] for i in range(30)])

    widths = {
        "A": 16, "B": 22, "C": 38, "D": 42, "E": 24, "F": 28, "G": 36, "H": 34, "I": 28, "J": 30
    }
    for ws in wb.worksheets:
        style_sheet(ws, widths)
    for ws in ["数据资料包", "客户问题池", "竞品拆解表", "首批15条脚本卡"]:
        for row in range(2, wb[ws].max_row + 1):
            wb[ws].row_dimensions[row].height = 54
    if cover_dir and (cover_dir / "封面模板三件套总览.png").exists():
        ws = wb["封面模板"]
        image = XLImage(str(cover_dir / "封面模板三件套总览.png"))
        image.width = 500
        image.height = 230
        ws.add_image(image, "A7")
        ws.row_dimensions[7].height = 175
    wb.save(out_path)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run.font.color.rgb = RGBColor(23, 54, 93)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9)


def build_doc(region: str, owner_role: str, account_line: str, out_path: Path, cover_dir: Path | None = None) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_heading(f"{region}大区总监自媒体口播落地执行手册", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("视频号优先｜研究先行｜15条可拍脚本｜封面正文｜拍摄发布复盘")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "一、账号定位与视频号设置")
    doc.add_paragraph(f"账号定位：{region}买房决策参谋。{owner_role}不做单盘复读，而做板块判断者 + 买房路径规划师。")
    add_table(doc, ["模块", "设置建议", "执行要求"], [
        ["账号名称", account_line, "用户指定账号名时逐字保留；不堆未经确认头衔。"],
        ["头像", "本人职业半身照", "背景干净，表情自然，避免项目效果图。"],
        ["简介", f"服务{region}改善家庭｜讲板块、预算、通勤和换房路径｜信息以官方与实际口径为准", "不放手机号、微信号等违规引流信息。"],
        ["置顶3条", f"{region}现在怎么看 / 预算怎么选 / 买前检查清单", "建立板块心智和咨询入口。"],
    ])

    add_heading(doc, "二、内容栏目与日更机制")
    add_table(doc, ["栏目", "占比", "作用", "示例选题"], [
        ["板块深挖", "40%", f"讲清{region}阶段、片区、兑现度", f"{region}不是过时，而是进入筛选期"],
        ["客户问题", "20%", "回应预算、置换、新房二手疑问", "预算220万怎么选"],
        ["配套交通", "15%", "用外景证明商业、轨道、公园是否好用", "地铁别只看近"],
        ["价格流通", "15%", "讲参考价、挂牌结构和未来转手逻辑", "挂牌多怎么看"],
        ["板块对比/案例", "10%", "对比礼嘉、悦来、照母山等替代板块", "两个板块怎么选"],
    ])
    doc.add_paragraph("执行节奏：每半个月集中拍摄1天，每次15条。拍摄前7天完成数据核验，前3天完成脚本，前1天完成通告单；拍摄后7天内15条全部入库。")

    add_heading(doc, "三、视频号封面、标题与正文规范")
    doc.add_paragraph("封面调用 cq-koubo-fengmian：默认生成数据判断型、客户问题型、实地验证/板块对比型三类封面。标题必须突出区域名，底部账号行固定或按用户指定。")
    if cover_dir and (cover_dir / "封面模板三件套总览.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(cover_dir / "封面模板三件套总览.png"), width=Inches(6.4))

    add_heading(doc, "四、首批15条可拍脚本（初稿）")
    for item in topics(region)[:15]:
        add_heading(doc, f"{item['day']}. {item['title']}", level=2)
        add_table(doc, ["项目", "内容"], [
            ["栏目/场景", f"{item['column']}｜{item['scene']}"],
            ["证据锚点", item["evidence"]],
            ["口播脚本", script_for_topic(region, item)],
            ["视频号正文", f"{region}买房不要只看一句结论，要看预算、通勤、产品和未来流通。评论区留预算和通勤方向。"],
            ["合规提示", "平台参考价、挂牌价、成交价分开表达；政策和规划以官方当期口径为准。"],
        ])

    add_heading(doc, "五、拍摄流程与现场执行")
    add_table(doc, ["时间", "动作", "交付物"], [
        ["拍摄前7天", "更新数据、客户问题、竞品拆解，锁定15条题目", "15条选题表"],
        ["拍摄前3天", "完成15条脚本，区域负责人确认口径", "脚本卡定稿"],
        ["拍摄前1天", "场地、服装、道具、提词器、拍摄顺序确认", "拍摄通告单"],
        ["拍摄日", "集中拍摄15条", "素材备份"],
        ["拍摄后7天", "15条成片全部入库", "发布库存"],
    ])

    add_heading(doc, "六、发布与复盘")
    add_table(doc, ["环节", "要求", "检查点"], [
        ["发布前", "标题、封面、正文、字幕、数据口径全部复核", "参考价、挂牌价、成交价分开说"],
        ["发布时", "优先测试12:30和20:00", "连续两周比较曝光和咨询"],
        ["发布后1小时", "回复评论、标记预算和板块问题", "不引导违规私下沟通"],
        ["每周复盘", "保留前20%题材，复拍相近问题", "形成下周选题"],
    ])

    add_heading(doc, "七、Excel执行表使用说明")
    doc.add_paragraph("配套Excel包含总览、数据资料包、客户问题池、竞品拆解表、封面模板、30天排期、15条脚本卡、拍摄通告单和发布复盘表。Excel是进度表和资料库，Word用于阅读和培训。")
    doc.save(out_path)


def generate_covers(region: str, account_line: str, output_dir: Path) -> Path | None:
    skills_dir = Path(__file__).resolve().parents[2]
    script = skills_dir / "cq-koubo-fengmian" / "scripts" / "create_cover_templates.py"
    if not script.exists():
        script = Path.home() / ".codex" / "skills" / "cq-koubo-fengmian" / "scripts" / "create_cover_templates.py"
    if not script.exists():
        return None
    cover_dir = output_dir / "封面模板"
    cmd = [
        sys.executable,
        str(script),
        "--region",
        region,
        "--account-line",
        account_line,
        "--output-dir",
        str(cover_dir),
    ]
    subprocess.run(cmd, check=True)
    return cover_dir


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成重庆区域口播运营Word+Excel落地包初稿")
    parser.add_argument("--region", required=True, help="区域/板块名，如礼嘉")
    parser.add_argument("--account-line", help="账号名，默认 {区域}买房参谋｜XX")
    parser.add_argument("--owner-role", default="区域总监", help="账号对象")
    parser.add_argument("--output-dir", help="输出目录，默认 ./{区域}口播运营落地包")
    parser.add_argument("--generate-covers", action="store_true", help="调用 cq-koubo-fengmian 生成封面模板")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    region = safe_name(args.region)
    account_line = args.account_line or f"{region}买房参谋｜XX"
    check_risk(region, account_line, args.owner_role)
    output_dir = Path(args.output_dir or f"{region}口播运营落地包").expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    cover_dir = generate_covers(region, account_line, output_dir) if args.generate_covers else None
    doc_path = output_dir / f"{region}大区总监自媒体口播落地执行手册.docx"
    xlsx_path = output_dir / f"{region}总监号运营执行表.xlsx"
    build_doc(region, args.owner_role, account_line, doc_path, cover_dir)
    build_workbook(region, account_line, xlsx_path, cover_dir)
    print(doc_path)
    print(xlsx_path)
    if cover_dir:
        print(cover_dir)


if __name__ == "__main__":
    main()
